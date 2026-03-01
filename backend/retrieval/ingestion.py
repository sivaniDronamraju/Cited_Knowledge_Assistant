# backend/retrieval/ingestion.py

import os
import json
import uuid
from typing import List, Union

import pandas as pd
from pypdf import PdfReader
import docx
from email import policy
from email.parser import BytesParser
from langdetect import detect, LangDetectException

from backend.retrieval.schemas import Document


 
# Metadata Utilities
 


def base_metadata(path: str, file_type: str) -> dict:
    """
    Build standardized metadata dictionary for a file.
    """
    return {
        "source_path": path,
        "file_name": os.path.basename(path),
        "file_type": file_type,
        "parent_folder": os.path.basename(os.path.dirname(path)),
    }


def detect_language(text: str) -> str:
    """
    Detect language of text safely.
    """
    try:
        return detect(text)
    except LangDetectException:
        return "unknown"


 
# Structured Data Handling
 


def row_to_sentence(row: pd.Series) -> str:
    """
    Convert structured tabular row into readable sentence.
    """

    parts = []

    for col, value in row.items():
        if pd.isna(value):
            continue

        col_clean = col.replace("_", " ").strip().title()

        if isinstance(value, str):
            if value.upper() == "Y":
                value = "Yes"
            elif value.upper() == "N":
                value = "No"

        parts.append(f"{col_clean}: {value}")

    if not parts:
        return ""

    return ". ".join(parts) + "."


 
# Parsers
 


def parse_docx(path: str) -> Document:
    doc = docx.Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)

    metadata = base_metadata(path, "docx")
    metadata["language"] = detect_language(text)

    return Document(
        document_id=str(uuid.uuid4()),
        text=text,
        metadata=metadata,
    )


def parse_pdf(path: str) -> Document:
    reader = PdfReader(path)
    text = ""

    for page in reader.pages:
        text += page.extract_text() or ""

    metadata = base_metadata(path, "pdf")
    metadata["language"] = detect_language(text)

    return Document(
        document_id=str(uuid.uuid4()),
        text=text,
        metadata=metadata,
    )


def parse_csv(path: str) -> List[Document]:
    df = pd.read_csv(path)
    documents: List[Document] = []

    for _, row in df.iterrows():
        sentence = row_to_sentence(row)
        if not sentence:
            continue

        metadata = base_metadata(path, "csv")
        metadata["language"] = "structured"

        documents.append(
            Document(
                document_id=str(uuid.uuid4()),
                text=sentence,
                metadata=metadata,
            )
        )

    return documents


def parse_xlsx(path: str) -> List[Document]:
    df = pd.read_excel(path)
    documents: List[Document] = []

    for _, row in df.iterrows():
        sentence = row_to_sentence(row)
        if not sentence:
            continue

        metadata = base_metadata(path, "xlsx")
        metadata["language"] = "structured"

        documents.append(
            Document(
                document_id=str(uuid.uuid4()),
                text=sentence,
                metadata=metadata,
            )
        )

    return documents


def parse_json(path: str) -> Document:
    text = ""

    if path.endswith(".jsonl"):
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                text += line.strip() + "\n"
    else:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            text = json.dumps(data, indent=2)

    metadata = base_metadata(path, "json")
    metadata["language"] = detect_language(text)

    return Document(
        document_id=str(uuid.uuid4()),
        text=text,
        metadata=metadata,
    )


def parse_eml(path: str) -> Document:
    with open(path, "rb") as f:
        msg = BytesParser(policy=policy.default).parse(f)

    subject = msg.get("subject", "")
    sender = msg.get("from", "")
    recipients = msg.get("to", "")

    body = ""

    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                body += part.get_content()
    else:
        body = msg.get_content()

    full_text = (
        f"Subject: {subject}\n"
        f"From: {sender}\n"
        f"To: {recipients}\n\n"
        f"{body}"
    )

    metadata = base_metadata(path, "eml")
    metadata["language"] = detect_language(full_text)

    return Document(
        document_id=str(uuid.uuid4()),
        text=full_text,
        metadata=metadata,
    )


def parse_txt(path: str) -> Document:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    metadata = base_metadata(path, "txt")
    metadata["language"] = detect_language(text)

    return Document(
        document_id=str(uuid.uuid4()),
        text=text,
        metadata=metadata,
    )


 
# Dispatcher
 


PARSERS = {
    ".docx": parse_docx,
    ".pdf": parse_pdf,
    ".csv": parse_csv,
    ".xlsx": parse_xlsx,
    ".json": parse_json,
    ".jsonl": parse_json,
    ".eml": parse_eml,
    ".txt": parse_txt,
}


def parse_file(path: str) -> Union[Document, List[Document], None]:
    ext = os.path.splitext(path)[1].lower()
    parser = PARSERS.get(ext)

    if parser is None:
        return None

    return parser(path)


def load_documents(root_path: str) -> List[Document]:
    """
    Recursively load and parse supported documents.
    """

    documents: List[Document] = []

    for root, _, files in os.walk(root_path):
        for file in files:
            path = os.path.join(root, file)
            parsed = parse_file(path)

            if parsed is None:
                continue

            if isinstance(parsed, Document):
                documents.append(parsed)
            elif isinstance(parsed, list):
                documents.extend(parsed)

    return documents