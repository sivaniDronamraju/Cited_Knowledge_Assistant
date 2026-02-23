# backend/ingestion.py

import os
import json
import uuid
import pandas as pd
from pypdf import PdfReader
import docx
from email import policy
from email.parser import BytesParser
from langdetect import detect, LangDetectException

from backend.schemas import Document


# -------------------------------------------------------
# Utility: Metadata Builder
# -------------------------------------------------------

def base_metadata(path: str, file_type: str):
    return {
        "source_path": path,
        "file_name": os.path.basename(path),
        "file_type": file_type,
        "parent_folder": os.path.basename(os.path.dirname(path)),
    }


# -------------------------------------------------------
# Utility: Language Detection
# -------------------------------------------------------

def detect_language(text: str) -> str:
    try:
        return detect(text)
    except LangDetectException:
        return "unknown"


# -------------------------------------------------------
# Utility: Structured Row → Sentence Conversion
# -------------------------------------------------------

def row_to_sentence(row: pd.Series) -> str:
    """
    Converts a structured row into a generic, readable sentence.

    - Skips null values
    - Converts underscores to spaces
    - Title-cases column names
    - Converts Y/N to Yes/No
    """

    parts = []

    for col, value in row.items():

        if pd.isna(value):
            continue

        # Clean column name
        col_clean = col.replace("_", " ").strip().title()

        # Normalize Yes/No flags
        if isinstance(value, str):
            if value.upper() == "Y":
                value = "Yes"
            elif value.upper() == "N":
                value = "No"

        parts.append(f"{col_clean}: {value}")

    return ". ".join(parts) + "."


# -------------------------------------------------------
# Parsers
# -------------------------------------------------------

def parse_docx(path: str):
    doc = docx.Document(path)
    text = "\n".join([p.text for p in doc.paragraphs])

    metadata = base_metadata(path, "docx")
    metadata["language"] = detect_language(text)

    return Document(
        document_id=str(uuid.uuid4()),
        text=text,
        metadata=metadata
    )


def parse_pdf(path: str):
    reader = PdfReader(path)
    text = ""

    for page in reader.pages:
        text += page.extract_text() or ""

    metadata = base_metadata(path, "pdf")
    metadata["language"] = detect_language(text)

    return Document(
        document_id=str(uuid.uuid4()),
        text=text,
        metadata=metadata
    )


def parse_csv(path: str):
    df = pd.read_csv(path)

    documents = []

    for _, row in df.iterrows():
        sentence = row_to_sentence(row)

        metadata = base_metadata(path, "csv")
        metadata["language"] = "structured"

        documents.append(
            Document(
                document_id=str(uuid.uuid4()),
                text=sentence,
                metadata=metadata
            )
        )

    return documents


def parse_xlsx(path: str):
    df = pd.read_excel(path)

    documents = []

    for _, row in df.iterrows():
        sentence = row_to_sentence(row)

        metadata = base_metadata(path, "xlsx")
        metadata["language"] = "structured"

        documents.append(
            Document(
                document_id=str(uuid.uuid4()),
                text=sentence,
                metadata=metadata
            )
        )

    return documents


def parse_json(path: str):
    text = ""

    if path.endswith(".jsonl"):
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                text += line.strip() + "\n"
    else:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            text = json.dumps(data, indent=2)

    metadata = base_metadata(path, "json")
    metadata["language"] = detect_language(text)

    return Document(
        document_id=str(uuid.uuid4()),
        text=text,
        metadata=metadata
    )


def parse_eml(path: str):
    with open(path, "rb") as f:
        msg = BytesParser(policy=policy.default).parse(f)

    subject = msg.get("subject", "")
    sender = msg.get("from", "")
    recipients = msg.get("to", "")

    body = ""
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                body += part.get_content()
    else:
        body = msg.get_content()

    full_text = f"Subject: {subject}\nFrom: {sender}\nTo: {recipients}\n\n{body}"

    metadata = base_metadata(path, "eml")
    metadata["language"] = detect_language(full_text)

    return Document(
        document_id=str(uuid.uuid4()),
        text=full_text,
        metadata=metadata
    )


def parse_txt(path: str):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    metadata = base_metadata(path, "txt")
    metadata["language"] = detect_language(text)

    return Document(
        document_id=str(uuid.uuid4()),
        text=text,
        metadata=metadata
    )


# -------------------------------------------------------
# Dispatcher
# -------------------------------------------------------

PARSERS = {
    ".docx": parse_docx,
    ".pdf": parse_pdf,
    ".csv": parse_csv,
    ".xlsx": parse_xlsx,
    ".json": parse_json,
    ".jsonl": parse_json,
    ".eml": parse_eml,
    ".txt": parse_txt,
}


def parse_file(path: str):
    for ext, parser in PARSERS.items():
        if path.endswith(ext):
            return parser(path)
    return None


def load_documents(root_path: str):
    documents = []

    for root, _, files in os.walk(root_path):
        for file in files:
            path = os.path.join(root, file)
            parsed = parse_file(path)

            if parsed is None:
                continue

            # Handle single Document
            if isinstance(parsed, Document):
                documents.append(parsed)

            # Handle list of Documents (CSV/XLSX)
            elif isinstance(parsed, list):
                documents.extend(parsed)

    return documents